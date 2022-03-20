import docx
from docx.shared import Pt, Inches
from docx.templates import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import ns
#from docx.styles import *
#from docx.image import *
from docx.text import *
from docxtpl import DocxTemplate, InlineImage
#from docx.enum.text import *


def starter_page(doc):
    """
    Add heading to page ==> None
    """
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header.paragraphs[0].text = r"Top Secret"


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    """
    Add page numbers to footer ==> None
    """
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def heading(head,doc):
    """
    Sets title of iamge ==> None
    """
    head.style = doc.styles['Heading 9']
    head.add_run().bold = True
    font = head.style.font
    font.name = 'Calibri'
    font.size = Pt(24)

def paragraph(p,doc):
    """
    Set paragraph settings ==> None
    """
    p.style = doc.styles['Body Text']
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = p.style.font
    form = p.style.paragraph_format
    form.keep_together = True
    font.name = 'David'
    font.size = Pt(12)

def build_doc(pic_dic):
    """
    Gets dict {'image-name':['image-path',text]} ==> doc obj
    """
    doc = word_obj()
    # Add footer
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number(doc.sections[0].footer.paragraphs[0].add_run())
    # Started Page
    starter_page(doc)
    for pic_name, pic_details in pic_dic.items():
        # Add head
        head = doc.add_heading(pic_name, 6)
        # add header
        heading(head,doc)
        # Add pic
        doc.add_picture(pic_details[0], width=Inches(6) ,height=Inches(4))
        # Add paragraph
        p = doc.add_paragraph(pic_details[1])
        paragraph(p,doc)
        ##
        doc.add_page_break()
    return doc

def word_obj():
    """
    Create word object ==> doc object
    """
    doc = docx.Document()
    return doc

def save(doc):
    """
    Saves doc ==> None
    """
    doc.save('Report/demo.docx')

def init():
    """
    test main
    """
    doc = build_doc({r"windows" : [r"..\files\report_images\win photo.png",r"this is the new windows"]})
    save(doc)
if __name__ == "__main__":
    init()